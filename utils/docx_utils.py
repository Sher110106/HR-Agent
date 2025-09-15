"""Utilities for generating DOCX documents.

Exposes helpers that return bytes ready for download buttons:
- text_to_docx(text, title)
- dataframe_to_docx_table(df, title)
- analysis_to_docx(analysis_text, dataframe, title)
"""

from __future__ import annotations

from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt


def _apply_default_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def _set_default_font(document: Document, size_pt: int = 11) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size_pt)


def _add_title(document: Document, title: Optional[str]) -> None:
    if not title:
        return
    document.add_heading(title, level=1)


def _add_inline_markdown_runs(paragraph, text: str) -> None:
    """Render a subset of Markdown inline formatting into runs.

    Supports bold **text** or __text__ and italics *text* or _text_.
    """
    i = 0
    n = len(text)
    bold = False
    italic = False
    buffer = []

    def flush_buffer(make_bold: bool, make_italic: bool) -> None:
        if not buffer:
            return
        run = paragraph.add_run("".join(buffer))
        run.bold = make_bold
        run.italic = make_italic
        buffer.clear()

    while i < n:
        # Handle bold markers ** or __
        if i + 1 < n and text[i : i + 2] in ("**", "__"):
            flush_buffer(bold, italic)
            bold = not bold
            i += 2
            continue
        # Handle italics markers * or _
        if text[i] in ("*", "_"):
            flush_buffer(bold, italic)
            italic = not italic
            i += 1
            continue
        buffer.append(text[i])
        i += 1

    flush_buffer(bold, italic)


def _add_markdown_content(document: Document, text: str) -> None:
    """Render simple Markdown into the document.

    Supported:
    - Headings: #, ##, ### (and a whole-line **Heading** treated as level-2)
    - Bullet lists: lines starting with '-' or '*'
    - Inline: bold (** or __) and italics (* or _)
    - Paragraphs separated by blank lines
    """
    lines = text.split("\n")
    paragraph_buffer: list[str] = []

    def flush_paragraph_buffer() -> None:
        if not paragraph_buffer:
            return
        para = document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        _add_inline_markdown_runs(para, "\n".join(paragraph_buffer))
        paragraph_buffer.clear()

    for raw in lines:
        line = raw.rstrip()
        if line.strip() == "":
            flush_paragraph_buffer()
            document.add_paragraph("")
            continue

        # Headings with #
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip('#')), 3)
            heading_text = line[level:].strip()
            if heading_text:
                document.add_heading(heading_text, level=level)
                continue

        # Heading if the whole line is bold like **Heading**
        if line.startswith("**") and line.endswith("**") and len(line) > 4 and line.count("**") == 2:
            heading_text = line[2:-2].strip()
            if heading_text:
                document.add_heading(heading_text, level=2)
                continue

        # Bullet list items
        if line.lstrip().startswith(("- ", "* ")):
            flush_paragraph_buffer()
            content = line.lstrip()[2:].strip()
            para = document.add_paragraph(style="List Bullet")
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            _add_inline_markdown_runs(para, content)
            continue

        # Normal paragraph content; accumulate until blank line
        paragraph_buffer.append(line)

    flush_paragraph_buffer()


def _bold_header_row(table) -> None:
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _enable_table_autofit(table) -> None:
    # Ensure table autofit is enabled in the underlying XML so Word adjusts widths
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.tblLayout
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "autofit")


def _document_to_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def text_to_docx(text: str, title: Optional[str] = None) -> bytes:
    """Create a DOCX containing the provided text, returned as bytes.

    - Title is optional and rendered as Heading 1 when provided.
    - Text preserves paragraphs (double newline) and line breaks (single newline).
    """
    doc = Document()
    _apply_default_page_margins(doc)
    _set_default_font(doc, size_pt=11)
    _add_title(doc, title)
    if text:
        _add_markdown_content(doc, text)
    return _document_to_bytes(doc)


def dataframe_to_docx_table(df, title: Optional[str] = None) -> bytes:  # type: ignore[no-untyped-def]
    """Create a DOCX with an optional title and a table for the DataFrame.

    Returns bytes suitable for Streamlit's download_button.
    """
    doc = Document()
    _apply_default_page_margins(doc)
    _set_default_font(doc, size_pt=11)
    _add_title(doc, title)

    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return _document_to_bytes(doc)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light List"
    _enable_table_autofit(table)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
    _bold_header_row(table)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = "" if value is None else str(value)

    return _document_to_bytes(doc)


def analysis_to_docx(
    analysis_text: Optional[str] = None,
    dataframe=None,  # type: ignore[no-untyped-def]
    title: Optional[str] = "Analysis Report",
) -> bytes:
    """Create a comprehensive report with analysis text and optional data table.

    - If analysis_text is provided, it is added as formatted paragraphs.
    - If dataframe is provided and non-empty, it is appended as a table.
    Returns bytes for direct download.
    """
    doc = Document()
    _apply_default_page_margins(doc)
    _set_default_font(doc, size_pt=11)
    _add_title(doc, title)

    if analysis_text:
        _add_markdown_content(doc, analysis_text)

    if dataframe is not None and getattr(dataframe, "empty", False) is False:
        doc.add_paragraph("")
        doc.add_heading("Data Table", level=2)
        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = "Light List"
        _enable_table_autofit(table)

        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(dataframe.columns):
            hdr_cells[i].text = str(col_name)
        _bold_header_row(table)

        for _, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = "" if value is None else str(value)

    return _document_to_bytes(doc)


